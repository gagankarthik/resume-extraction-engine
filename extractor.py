"""
Resume text extraction — no OCR dependency.
  PDF  → pdfplumber (layout-aware, table-aware, repeated-header removal)
  DOCX → python-docx (paragraphs + tables + headers/footers in doc order)
           fallback: raw ZIP extraction of word/document.xml
  RTF  → striprtf (control words stripped, escapes and unicode decoded)
  DOC  → not parseable; explains how to convert. Files that merely *claim*
         .doc are usually RTF or DOCX and are routed by their real bytes,
         so this message is reached only by genuine OLE2 documents.
  TXT  → UTF-8 decode
All paths run through normalizer.py before returning.
"""
import io
import math
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from striprtf.striprtf import rtf_to_text

from normalizer import deduplicate_page_content, normalize_text

# Characters of text below which a PDF page is considered unparseable
_SPARSE_THRESHOLD = 50

ExtractionResult = tuple[str, int, dict]


# ------------------------------------------------------------------ #
# Column detection
# ------------------------------------------------------------------ #
#
# A two-column resume read as one column is not slightly wrong, it is unusable.
# pdfplumber's layout mode pads text to its x-position, so a sidebar line and a
# body line share one output line:
#
#     Kubernetes                          • Built the ingestion pipeline
#
# The normalizer then collapses the padding to a single space and the skill is
# welded onto the front of somebody's job bullet. Worse, the bullet glyph is no
# longer at the start of the line, so the programmatic bullet count comes back 0
# — and a job with zero expected bullets is skipped by the validator entirely.
# The accuracy checks switch themselves off on exactly the layouts that need
# them most, and nothing in the output says so.
#
# The fix is to find the gutter and read each column on its own.

# Narrower gaps than this are word spacing or a right-aligned date, not a gutter.
_MIN_GUTTER_POINTS = 14

# Each column must hold a real share of the page. Without this, the whitespace
# before a right-aligned date column ("Acme Corp        2019 - Present") reads as
# a gutter and the page is split down a seam that isn't there.
_MIN_COLUMN_WORD_SHARE = 0.15

# ...and span enough lines to be a column rather than a heading that happens to
# sit alone on its side of the page.
_MIN_COLUMN_LINES = 4

# A gutter near the edge of the text block is a margin artefact.
_INTERIOR_MARGIN = 0.15


def detect_column_bands(words: list[dict]) -> list[tuple[float, float]] | None:
    """The x-ranges of the page's columns, or None if it reads as one column.

    Works from word boxes rather than rendered text: a gutter is a vertical
    strip that no word on the page crosses. That test is what keeps ordinary
    single-column resumes intact — a bullet line runs the full width of the text
    block, so it crosses any candidate seam and rules it out, while the gap in
    front of a right-aligned date does not extend down the page.

    `words` are pdfplumber word dicts (x0, x1, top). Kept pure and separate from
    the page object so the rule can be tested without rendering a PDF.
    """
    if len(words) < 20:
        return None

    left = min(w["x0"] for w in words)
    right = max(w["x1"] for w in words)
    width = right - left
    if width <= 0:
        return None

    # One bin per point, marked where any word sits.
    bins = math.ceil(width) + 1
    occupied = bytearray(bins)
    for w in words:
        start = max(0, int(w["x0"] - left))
        end = min(bins - 1, math.ceil(w["x1"] - left))
        for i in range(start, end + 1):
            occupied[i] = 1

    gutters: list[tuple[int, int]] = []
    i = 0
    while i < bins:
        if occupied[i]:
            i += 1
            continue
        j = i
        while j < bins and not occupied[j]:
            j += 1
        if j - i >= _MIN_GUTTER_POINTS:
            centre = (i + j) / 2
            if _INTERIOR_MARGIN * width < centre < (1 - _INTERIOR_MARGIN) * width:
                gutters.append((i, j))
        i = j

    if not gutters:
        return None

    bands: list[tuple[float, float]] = []
    cursor = left
    for start, end in gutters:
        bands.append((cursor, left + start))
        cursor = left + end
    bands.append((cursor, right))

    # Every band has to look like a column, or the split is an artefact and the
    # page is safer read whole.
    total = len(words)
    for x0, x1 in bands:
        in_band = [w for w in words if x0 <= (w["x0"] + w["x1"]) / 2 <= x1]
        if len(in_band) < _MIN_COLUMN_WORD_SHARE * total:
            return None
        if len({round(w["top"]) for w in in_band}) < _MIN_COLUMN_LINES:
            return None

    return bands


# ------------------------------------------------------------------ #
# Public entry point
# ------------------------------------------------------------------ #

def extract_text(file_bytes: bytes, file_type: str) -> ExtractionResult:
    """
    Returns (normalized_text, page_count, extraction_info).
    Raises ValueError for unsupported types.
    Raises RuntimeError for unreadable files, with a message meant for the
    person who uploaded them.

    `file_type` is filetypes.resolve()'s verdict — decided from the file's own
    bytes, not its name — so each branch here can trust what it is given.
    """
    handler = _HANDLERS.get(file_type.lower())
    if handler is None:
        raise ValueError(f"Unsupported file type: {file_type!r}")
    return handler(file_bytes)


# ------------------------------------------------------------------ #
# PDF
# ------------------------------------------------------------------ #

def _extract_pdf(file_bytes: bytes) -> ExtractionResult:
    # Imported here rather than at module scope: pdfplumber pulls in pdfminer
    # and Pillow, seconds of cold start on a Lambda that may well have been
    # handed a DOCX.
    import pdfplumber

    page_texts: list[str | None] = []
    sparse_pages: list[int] = []
    multi_column_pages: list[int] = []

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for i, page in enumerate(pdf.pages, start=1):
                text, columns = _extract_one_pdf_page(page)
                if columns > 1:
                    multi_column_pages.append(i)
                if len(text.strip()) < _SPARSE_THRESHOLD:
                    sparse_pages.append(i)
                page_texts.append(text or None)
    except Exception as exc:
        # A damaged or encrypted PDF is the uploader's problem to fix, not a
        # fault in this service — say which, and what to do, rather than 500.
        if "password" in str(exc).lower() or "encrypt" in str(exc).lower():
            raise RuntimeError(
                "This PDF is password-protected, so its text cannot be read. "
                "Remove the password, or export an unprotected copy, and upload that."
            ) from exc
        raise RuntimeError(
            "This PDF could not be opened — the file looks damaged or incomplete. "
            "Try re-downloading it, or open it and export a fresh PDF."
        ) from exc

    if sparse_pages and all(t is None or len(t.strip()) < _SPARSE_THRESHOLD for t in page_texts):
        raise RuntimeError(
            "This PDF appears to be a scanned image — no machine-readable text found. "
            "Please use a text-based PDF or a DOCX file."
        )

    # Remove running headers/footers that repeat across pages
    page_texts = deduplicate_page_content(page_texts)

    combined = "\n\n".join(t for t in page_texts if t)
    normalized = normalize_text(combined)

    return normalized, page_count, {
        "method": "pdfplumber",
        "sparse_pages": sparse_pages,
        # Reported so a layout the reader had to take apart is visible when
        # someone is working out why a particular resume came back oddly.
        "multi_column_pages": multi_column_pages,
    }


def _extract_one_pdf_page(page) -> tuple[str, int]:
    """
    Extract text from one pdfplumber Page, returning (text, column_count).

    Columns are read one at a time when the page has them. Reading a
    two-column page in one pass interleaves the sidebar with the body — see
    detect_column_bands for what that costs — so the gutter is found first and
    each side is cropped and read on its own, left to right.
    """
    columns = _read_columns(page)
    if columns is not None:
        text = "\n\n".join(columns)
        column_count = len(columns)
    else:
        # layout=True preserves spatial ordering within a single column
        text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3) or ""
        if not text.strip():
            text = page.extract_text() or ""
        column_count = 1

    # Explicit table extraction — pdfplumber may miss grid-based tables in layout mode
    table_rows: list[str] = []
    for table in page.extract_tables():
        for row in table:
            cells = [str(c).strip() for c in row if c and str(c).strip()]
            if cells:
                table_rows.append(" | ".join(cells))

    if table_rows:
        table_block = "\n".join(table_rows)
        # Only append if the content is not already present in the main text
        if table_block[:60] not in text:
            text = text + "\n\n" + table_block

    return text.strip(), column_count


def _read_columns(page) -> list[str] | None:
    """Each column's text in reading order, or None if the page has one column.

    Any failure here returns None, which falls back to reading the page whole —
    the behaviour that shipped before columns were detected at all. A layout
    this cannot make sense of should degrade to the old result, never to no text.
    """
    try:
        words = page.extract_words(x_tolerance=1.5, y_tolerance=3)
    except Exception:
        return None

    bands = detect_column_bands(words)
    if not bands:
        return None

    _, top, _, bottom = page.bbox
    texts: list[str] = []
    for x0, x1 in bands:
        try:
            # A point of margin either side, so a glyph sitting exactly on the
            # boundary is not clipped in half.
            crop = page.crop((
                max(page.bbox[0], x0 - 1),
                top,
                min(page.bbox[2], x1 + 1),
                bottom,
            ))
            text = (crop.extract_text(layout=True, x_tolerance=3, y_tolerance=3) or "").strip()
        except Exception:
            return None
        if text:
            texts.append(text)

    # One column's worth of text means the split found nothing worth having.
    return texts if len(texts) > 1 else None


# ------------------------------------------------------------------ #
# DOCX
# ------------------------------------------------------------------ #

# Legacy Word .doc files are OLE2 compound documents (magic D0 CF 11 E0 A1 B1 1A E1).
# Modern .docx files are ZIP archives (magic PK\x03\x04). python-docx only reads
# the latter, so detect the former up front and fail with an actionable message
# instead of letting python-docx + the ZIP fallback both throw opaque errors.
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

# Reading OLE2 needs a Word-format parser (antiword, LibreOffice) that cannot be
# shipped in a Lambda zip. Saying exactly which conversions work is the whole
# value of this message — "unsupported" alone leaves someone stuck.
_LEGACY_DOC_MESSAGE = (
    "This is a legacy Word document (.doc), a format from before 2007 that this "
    "tool cannot read. Open it in Word or Google Docs and use “Save as” to make "
    "a .docx, or export it to PDF — either will upload fine."
)


def _extract_doc(file_bytes: bytes) -> ExtractionResult:
    """
    Genuine OLE2 .doc.

    Most uploads named .doc never reach here: filetypes.resolve() reads their
    bytes first, and a "resume.doc" that is really RTF or DOCX is routed to the
    parser that can read it. What is left is the real thing, which cannot be.
    """
    raise RuntimeError(_LEGACY_DOC_MESSAGE)


def _extract_docx(file_bytes: bytes) -> ExtractionResult:
    # Defence in depth: routing should have sent this to _extract_doc already.
    if file_bytes[:8] == _OLE2_MAGIC:
        raise RuntimeError(_LEGACY_DOC_MESSAGE)

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as docx_exc:
        try:
            return _extract_docx_zip_fallback(file_bytes)
        except Exception:
            raise RuntimeError(
                f"Could not read this Word file ({docx_exc}). "
                "Try re-saving it as .docx in Microsoft Word, or convert to PDF."
            ) from docx_exc

    parts: list[str] = []

    # Section headers
    for section in doc.sections:
        for attr in ("header", "even_page_header", "first_page_header"):
            try:
                hdr = getattr(section, attr)
                for para in hdr.paragraphs:
                    t = para.text.strip()
                    if t:
                        parts.append(t)
            except Exception:
                pass

    # Body in document order: paragraphs then tables
    for block in _iter_docx_blocks(doc):
        if isinstance(block, DocxParagraph):
            text = block.text.strip()
            if not text:
                continue
            style = ((block.style.name if block.style else None) or "").lower()
            is_heading = "heading" in style or "title" in style or (
                text.isupper() and 2 < len(text) < 80
            )
            # Word stores bullet/numbered list items as paragraphs with a <w:numPr>
            # element but NO text glyph — the bullet is rendered from the numbering
            # definition. Without re-introducing a glyph here, the normalizer and
            # StructureAgent can't tell that "Designed and implemented X" was a
            # bullet vs a free-floating paragraph, and bullet_count comes out 0.
            # Prepend "• " so downstream regex bullet detection works.
            if not is_heading and _is_list_paragraph(block):
                text = f"• {text}"
            parts.append(("\n" + text) if is_heading else text)

        elif isinstance(block, DocxTable):
            for row in block.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

    # Section footers (skip pure page numbers)
    for section in doc.sections:
        for attr in ("footer", "even_page_footer", "first_page_footer"):
            try:
                ftr = getattr(section, attr)
                for para in ftr.paragraphs:
                    t = para.text.strip()
                    if t and not t.isdigit():
                        parts.append(t)
            except Exception:
                pass

    combined = "\n".join(parts)
    normalized = normalize_text(combined)
    return normalized, 1, {"method": "docx"}


def _extract_docx_zip_fallback(file_bytes: bytes) -> ExtractionResult:
    """Extract text from word/document.xml directly when python-docx rejects the file."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
        names = z.namelist()
        doc_xml = next(
            (n for n in names if n == "word/document.xml"),
            next((n for n in names if "document" in n.lower() and n.endswith(".xml")), None),
        )
        if doc_xml is None:
            raise RuntimeError("No document.xml found inside ZIP archive")
        data = z.read(doc_xml)

    root = ET.fromstring(data)
    parts: list[str] = []
    for para in root.iter(f"{{{W}}}p"):
        text = "".join(t.text or "" for t in para.iter(f"{{{W}}}t")).strip()
        if not text:
            continue
        # Mirror the python-docx path: prepend "• " for paragraphs in a list.
        pPr = para.find(f"{{{W}}}pPr")
        is_list = pPr is not None and pPr.find(f"{{{W}}}numPr") is not None
        if is_list:
            text = f"• {text}"
        parts.append(text)

    combined = "\n".join(parts)
    return normalize_text(combined), 1, {"method": "docx_zip_fallback"}


def _iter_docx_blocks(doc: Document):
    """Yield Paragraph and Table in document body order."""
    P_TAG = qn("w:p")
    T_TAG = qn("w:tbl")
    for child in doc.element.body:
        if child.tag == P_TAG:
            yield DocxParagraph(child, doc)
        elif child.tag == T_TAG:
            yield DocxTable(child, doc)


def _is_list_paragraph(para: DocxParagraph) -> bool:
    """True if this paragraph is part of a bulleted or numbered list in Word.

    Detection is restricted to the presence of <w:numPr>, which is Word's
    definitive list-membership signal. Style name alone (e.g. "ListParagraph")
    is unreliable — Word applies it to plain indented prose too.
    """
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


# ------------------------------------------------------------------ #
# Plain text
# ------------------------------------------------------------------ #

def _extract_txt(file_bytes: bytes) -> ExtractionResult:
    text = _decode_text(file_bytes)
    return normalize_text(text), 1, {"method": "txt"}


def _decode_text(file_bytes: bytes) -> str:
    """
    Decode bytes that are meant to be text.

    Resumes routinely arrive as Windows-1252 from Word or Notepad — smart
    quotes, em dashes and accented names are exactly the bytes UTF-8 rejects.
    Falling straight to errors="replace" would pepper those names with U+FFFD,
    so try the encodings that actually occur before giving up on any character.
    """
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


# ------------------------------------------------------------------ #
# RTF
# ------------------------------------------------------------------ #

def _extract_rtf(file_bytes: bytes) -> ExtractionResult:
    """
    RTF is text with markup, so it decodes before it parses.

    striprtf drops control words, skips destination groups such as font and
    colour tables, and turns \\'xx and \\uN escapes back into characters —
    including the \\bullet that list items are built from, which downstream
    bullet counting depends on.
    """
    try:
        text = rtf_to_text(_decode_text(file_bytes), errors="ignore")
    except Exception as exc:
        raise RuntimeError(
            "This RTF file could not be read — it may be damaged. Try opening it "
            "in Word or TextEdit and saving it again as .docx or PDF."
        ) from exc

    if not text.strip():
        raise RuntimeError(
            "This RTF file contains no readable text. If the resume is a picture "
            "pasted into the document, upload the original instead — text has to "
            "be selectable, not pictured."
        )

    return normalize_text(text), 1, {"method": "striprtf"}


# ------------------------------------------------------------------ #
# Dispatch
# ------------------------------------------------------------------ #
# Keys are FileKind.key values from filetypes.py. Adding a format means adding
# a FileKind there and one entry here; nothing else in the service changes.

_HANDLERS = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "doc": _extract_doc,
    "rtf": _extract_rtf,
    "txt": _extract_txt,
}
